import os
import re
import pandas as pd
import pdfplumber
from docx import Document
from docx.enum.section import WD_ORIENT
from collections import Counter
import string
from datetime import datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

# legacy defaults (can remain unset when imported from other code)
PDF_PATH = None
EXCEL_PATH = None

TEMPLATE_PATH = r"C:\NCRP\Sample_Updated.docx"
IFSC_CSV_PATH = r"C:\NCRP\IFSC_CODES_FOR_ANALYSIS.csv"
OUTPUT_DIR = r"C:\NCRP\OUT2"

# The module provides a legacy script that reads a PDF and an excel file to
# generate bank-wise Word letters.  The original `generate_letters()` entry
# point used hardcoded global constants; the refactor below exposes a
# reusable function that accepts paths so the Flask backend can call it with
# runtime parameters.

IFSC_DICT = {}
def load_ifsc_csv(path):
    global IFSC_DICT
    df = pd.read_csv(path, dtype=str).fillna('')
    df.columns = df.columns.str.strip().str.upper()
    ifsc_col = next((c for c in df.columns if 'IFSC' in c), None)
    bank_col = next((c for c in df.columns if 'BANK' in c), None)
    for _, r in df.iterrows():
        prefix = str(r[ifsc_col]).strip().upper()[:4]
        bank = str(r[bank_col]).strip().upper()
        IFSC_DICT[prefix] = bank

def extract_pdf_text():
    text = ''
    with pdfplumber.open(PDF_PATH) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                text += t + '\n'
    return text

def extract_pdf_placeholders(text):
    data = {}

    def find(pattern, default=None):
        m = re.search(pattern, text, re.I | re.DOTALL)
        return m.group(1).strip() if m else default

    ncrp_raw = find(r'Acknowledgement\s*No\.?\s*[:\-]?\s*([\d]{5,20})')
    data['{{NCRPNO}}'] = ncrp_raw if ncrp_raw else 'N/A'

    def find(pattern, default=None, group=1):
        m = re.search(pattern, text, re.I | re.DOTALL)
        return m.group(group).strip() if m else default

    csr = find(r'(FIR|Crime|CSR|Cr\.?)\s*No[:\-\s]*([A-Za-z0-9/ \-]+)', group=2)
    data['{{CSRNO}}'] = csr if csr else 'N/A'

    complaint = find(r'Complaint Additional Info\s*(.*?)\s*(Total Fraudulent Amount|$)')
    data['{{Complaint Additional Info}}'] = complaint if complaint else 'N/A'

    name = find(r'Complainant\s*Name[:\s]+([A-Za-z .\']{1,100})')
    data['{{COM_NAME}}'] = name if name else 'N/A'

    accounts = re.findall(r'\b\d{10,20}\b', text)
    data['{{COM_AC_NO}}'] = accounts[0] if accounts else 'N/A'

    banks = re.findall(r'\b([A-Z][A-Za-z ]+Bank)\b', text)
    data['{{COM_BANK}}'] = Counter(banks).most_common(1)[0][0] if banks else 'N/A'
    
    total_amount = find(r'Total\s*Fraudulent\s*Amount\s*reported\s*by\s*complainant\s*[-:]?\s*[:\-]?\s*([\d,]+\.\d{2}|\d[\d,]*)')
    data['{{TOTAL_FRAUD_AMOUNT}}'] = total_amount if total_amount else 'N/A'

    return data

def get_full_bank_name(ifsc_code):
    prefix = str(ifsc_code)[:4].upper()
    bank = IFSC_DICT.get(prefix)
    return bank.title() if bank else (prefix + " BANK").title()

def strong_replace(doc, map_):

    def replace_in_paragraph(p):

        full_text = "".join([r.text for r in p.runs])
        replaced = full_text

        for k, v in map_.items():
            replaced = replaced.replace(k, str(v))

        if replaced != full_text:

            if p.runs:
                for i in range(len(p.runs) - 1, 0, -1):
                    p.runs[i].clear()

                complaint_value = map_.get("{{Complaint Additional Info}}", "")

                
                if complaint_value and complaint_value in replaced:

                    cleaned_text = re.sub(r'\s+', ' ', replaced).strip()
                    p.runs[0].text = cleaned_text

                    p.paragraph_format.left_indent = None
                    p.paragraph_format.right_indent = None
                    p.paragraph_format.first_line_indent = None

                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.paragraph_format.line_spacing = 1.5
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(6)

                else:
                    
                    p.runs[0].text = replaced

    for p in doc.paragraphs:
        replace_in_paragraph(p)

    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p)

    for section in doc.sections:
        for p in section.header.paragraphs:
            replace_in_paragraph(p)

    for section in doc.sections:
        for p in section.footer.paragraphs:
            replace_in_paragraph(p)

    return doc

def insert_rows(doc, rows):
    for t in doc.tables:
        if "S.NO" in t.rows[0].cells[0].text.upper():
            for r in rows:
                cells = t.add_row().cells
                for i in range(min(len(r), len(cells))):
                    cells[i].text = str(r[i])
            return True
    return False

def clean_fn(s):
    allowed = "-_.()%s%s" % (string.ascii_letters, string.digits)
    return ''.join(c for c in s if c in allowed).replace(" ", "_")


def _read_excel_or_csv(path):
    """Return a pandas DataFrame from either an Excel or CSV file."""
    if path.lower().endswith('.csv'):
        return pd.read_csv(path, dtype=str).fillna('')
    else:
        return pd.read_excel(path, dtype=str).fillna('')


def generate_letters_from_files(pdf_path,
                                excel_path,
                                output_dir=None,
                                template_path=None,
                                ifsc_csv_path=None):
    """Generate the bank-wise letters and return a list of generated files.

    Arguments are mostly self‑explanatory; any parameter left as None will fall
    back to the module global constants.  Raised exceptions propagate up to the
    caller to handle (for example the Flask route)."""
    tpl = template_path or TEMPLATE_PATH
    ifsc_csv = ifsc_csv_path or IFSC_CSV_PATH
    out_dir = output_dir or OUTPUT_DIR

    # load IFSC reference
    load_ifsc_csv(ifsc_csv)

    # read the PDF text and placeholders
    pdf_text = ''
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                pdf_text += t + '\n'
    pdf_data = extract_pdf_placeholders(pdf_text)

    # read table data
    df = _read_excel_or_csv(excel_path)

    df_final = pd.DataFrame({
        'S.NO': range(1, len(df) + 1),
        'LAYER': df.iloc[:, 5].astype(str),
        'SUSPECT/BENEFICIARY DETAILS': df.iloc[:, 6].astype(str),
        'SUSPECT/IFSC_CODE': df.iloc[:, 7].astype(str),
        'TXN_ID / UTR_NO': df.iloc[:, 9].astype(str),
        'DISPUTED AMOUNT': df.iloc[:, 11].astype(str),
        'TXN AMOUNT': df.iloc[:, 10].astype(str).str.replace(',', '').str.strip()
    })

    df_final['SUSPECT/IFSC_CODE'] = (
        df_final['SUSPECT/IFSC_CODE']
        .astype(str)
        .str.strip()
        .str.upper()
    )
    df_final = df_final[df_final['SUSPECT/IFSC_CODE'] != '']
    df_final['BANK_CODE'] = df_final['SUSPECT/IFSC_CODE'].str[:4]

    generated_files = []

    for bank_code, group in df_final.groupby('BANK_CODE'):
        rows = []
        for idx, (_, r) in enumerate(group.iterrows(), 1):
            rows.append([
                idx,
                r['LAYER'],
                r['SUSPECT/BENEFICIARY DETAILS'],
                r['SUSPECT/IFSC_CODE'],
                r['TXN_ID / UTR_NO'],
                r['DISPUTED AMOUNT'],
                r['TXN AMOUNT']
            ])

        doc = Document(tpl)
        inserted = insert_rows(doc, rows)
        if not inserted:
            # if the template doesn't have the expected table, skip this bank
            continue

        sample_ifsc = group['SUSPECT/IFSC_CODE'].iloc[0]
        full_bank_name = get_full_bank_name(sample_ifsc)

        replacements = pdf_data.copy()
        replacements['{{BANK_NAME}}'] = full_bank_name
        replacements['{{GETDATE}}'] = datetime.now().strftime('%d-%m-%Y')

        doc = strong_replace(doc, replacements)

        csr_part = replacements.get('{{CSRNO}}', 'NA').replace('/', '-')
        filename = f"{clean_fn(full_bank_name)}__CSR_{clean_fn(csr_part)}.docx"

        os.makedirs(out_dir, exist_ok=True)
        out_path = os.path.join(out_dir, filename)
        doc.save(out_path)
        generated_files.append(out_path)

    return generated_files


def generate_letters():
    # backwards compatible entry point for CLI usage
    if not PDF_PATH or not EXCEL_PATH:
        print("ERROR: PDF_PATH and EXCEL_PATH must be set when running as a script")
        return
    generated = generate_letters_from_files(PDF_PATH, EXCEL_PATH, OUTPUT_DIR, TEMPLATE_PATH, IFSC_CSV_PATH)
    for f in generated:
        print(f"Generated: {f}")


if __name__ == "__main__":
    generate_letters()

if __name__ == "__main__":
    generate_letters()
